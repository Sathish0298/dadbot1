import streamlit as st
from io import StringIO
import requests
from docx import Document
from bs4 import BeautifulSoup
import pandas as pd
from urllib.parse import urlparse


# Function to determine the website based on the URL
def determine_website(url):
    parsed_url = urlparse(url)
    hostname = parsed_url.hostname

    if "amazon" in hostname:
        return "amazon"
    elif "flipkart" in hostname:
        print("Flipkart")
        return "flipkart"
    elif "snapdeal" in hostname:
        return "snapdeal"
    else:
        return None


# Amazon review scraping function
def amazon_review_scraper(url, page):
    url = f"{url}{page}"

    reviews = []

    user_agent = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/121.0.0.0 Safari/537.36 Edg/121.0.0.0"
    }

    response = requests.get(url=url, headers=user_agent)
    soup = BeautifulSoup(response.content, "lxml")

    for review in soup.find_all("div", {"class": "a-section review aok-relative"}):
        name = review.find("span", {"class": "a-profile-name"}).text
        rating = review.find("i", {"data-hook": "review-star-rating"}).text
        comments = review.find(
            "div", {"class": "a-row a-spacing-small review-data"}
        ).text

        data = {
            "Name": name,
            "Rating": rating,
            "Comments": comments,
        }

        reviews.append(data)

    return reviews


# Flipkart review scraping function
def flipkart_review_scraper(url, page):
    url = f"{url}{page}"

    reviews = []

    user_agent = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/121.0.0.0 Safari/537.36 Edg/121.0.0.0"
    }

    response = requests.get(url=url, headers=user_agent)
    soup = BeautifulSoup(response.content, "html.parser")

    for review in soup.find_all(
        "div", {"class": "_27M-vq"}
    ):  # Updated class for review container
        name = review.find(
            "p", {"class": "_2sc7ZR _2V5EHH"}
        ).text  # Updated class for reviewer's name
        rating = review.find(
            "div", {"class": "_3LWZlK"}
        ).text  # Updated class for rating
        comments = review.find(
            "div", {"class": "t-ZTKy"}
        ).text  # Updated class for review comments
        comments = re.sub(
            r"\s*READ\s+MORE\s*", "", comments
        )  # Remove 'READ MORE' links if present

        data = {
            "Name": name,
            "Rating": rating,
            "Comments": comments.strip(),
        }

        reviews.append(data)

    return reviews


# Snapdeal review scraping function
def snapdeal_review_scraper(url, page):
    url = f"{url}{page}"
    reviews = []

    user_agent = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/121.0.0.0 Safari/537.36 Edg/121.0.0.0"
    }

    response = requests.get(url=url, headers=user_agent)
    soup = BeautifulSoup(response.content, "html.parser")

    reviews_skipped = 0  # Counter to track the number of reviews skipped

    for review in soup.find_all("div", {"class": "user-review"}):
        if reviews_skipped < 2:
            reviews_skipped += 1
            continue  # Skip the first two reviews

        rating_elements = review.find_all("i", class_="sd-icon sd-icon-star active")
        rating = len(rating_elements)
        name = review.find("div", {"class": "_reviewUserName"}).get("title")
        comments = review.find("p").text

        data = {
            "Name": name,
            "Rating": rating,
            "Comments": comments,
        }

        reviews.append(data)

    return reviews


def main():
    st.set_page_config(page_title="DaD BOT", layout="wide")

    # Custom CSS to style the Streamlit app
    st.markdown(
        """
        <style>
        .big-font {
            font-size:30px !important;
        }
        .button-style {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 100%;
            height: 3rem;
            padding: .5rem;
            margin: .5rem 0;
            background-color: #0c6efc;
            color: red;
            border-radius: 5px;
            border: none;
            font-size: 18px;
        }
        .streamlit-input {
            margin-bottom: 20px;
        }
        .streamlit-header {
            margin-bottom: 5px;
        }
        </style>
    """,
        unsafe_allow_html=True,
    )

    # Header
    st.markdown('<p class="big-font">DaD BOT</p>', unsafe_allow_html=True)

    # Navigation bar
    with st.sidebar:
        st.title("Navigation")
        st.button("Home")
        st.button("Register")
        st.button("Login")

    with st.form("product_link_form"):
        product_link = st.text_input("Paste the link to the product:")
        submit_link = st.form_submit_button("Get the Product Review")
        if submit_link and product_link:
            initiate_scraping(product_link)

    # Display and Download Scraped Data Section
    st.header("View and Download Scraped Data")
    if st.button("View Scraped Data"):
        display_csv()
    download_csv()

    # Product Specification Section
    with st.container():
        st.markdown("## Display the new product specification", unsafe_allow_html=True)
        if st.button("View The Specification", key="product_spec"):
            view_product_specification()


def scrape_reviews(urls):
    url_suffix = ""
    scraper_func = None
    all_reviews = []
    for url in urls:
        website = determine_website(url)
        if not website:
            print(f"Unsupported website for URL: {url}")
            continue

        print(f"Scraping {website} reviews from {url}")

        if website == "amazon":
            scraper_func = amazon_review_scraper
            url_suffix = "?th=1&pageNumber="
        elif website == "flipkart":
            scraper_func = flipkart_review_scraper
            url_suffix = "&page="
        elif website == "snapdeal":
            scraper_func = snapdeal_review_scraper
            url_suffix = "?page="

        # Assuming we are scraping only the first 3 pages for demonstration
        for page in range(1, 2):
            full_url = f"{url}{url_suffix}{page}"
            reviews = scraper_func(full_url, page)
            all_reviews.extend(reviews)

    # Convert list of reviews to a DataFrame and save to CSV string
    reviews_df = pd.DataFrame(all_reviews)
    csv_string = reviews_df.to_csv(index=False)
    return csv_string


# Function to initiate scraping and save the results
def initiate_scraping(url):
    reviews_csv = scrape_reviews([url])
    st.session_state["reviews_csv"] = reviews_csv
    st.success("Scraped data has been updated!")


# Function to display the CSV content
def display_csv():
    if "reviews_csv" in st.session_state and st.session_state["reviews_csv"]:
        df = pd.read_csv(StringIO(st.session_state["reviews_csv"]))
        st.dataframe(df.head(10))
    else:
        st.error("No data available. Please scrape some data first.")


# Function to download the CSV file
def download_csv():
    if "reviews_csv" in st.session_state and st.session_state["reviews_csv"]:
        st.download_button(
            label="Download CSV",
            data=st.session_state["reviews_csv"],
            file_name="reviews.csv",
            mime="text/csv",
        )
    else:
        st.error("No data available to download.")


def view_product_specification():
    doc_path = "outputs/reviews_analyzed_full_improvements_prioritized.docx"

    try:
        doc = Document(doc_path)
        if doc.tables:
            st.markdown("## Top Improvement Suggestions")
            # Iterate over each table in the document
            for table in doc.tables:
                # Extract the content of the table
                suggestions = [
                    cell.text.strip() for row in table.rows for cell in row.cells
                ]

                # Filter out empty strings
                suggestions = [suggestion for suggestion in suggestions if suggestion]

                # Display each suggestion as a bullet point
                for suggestion in suggestions:
                    st.markdown(f"- {suggestion}")
        else:
            st.error("No tables found in the document.")
    except Exception as e:
        st.error(f"Failed to read the document: {str(e)}")


if __name__ == "__main__":
    main()
